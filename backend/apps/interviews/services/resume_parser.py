"""Parse resume PDF/text into structured profile hints."""

from __future__ import annotations

import re


def extract_text_from_upload(uploaded_file) -> str:
    """Best-effort text extraction from uploaded resume."""
    if not uploaded_file:
        return ""
    name = (getattr(uploaded_file, "name", "") or "").lower()
    raw = uploaded_file.read()
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)

    if name.endswith(".pdf"):
        try:
            import pypdf

            reader = pypdf.PdfReader(uploaded_file)
            parts = []
            for page in reader.pages[:12]:
                parts.append(page.extract_text() or "")
            return "\n".join(parts).strip()
        except Exception:
            try:
                uploaded_file.seek(0)
            except Exception:
                pass
            return raw.decode("utf-8", errors="ignore")[:20000]

    if name.endswith((".doc", ".docx")):
        try:
            import io

            from docx import Document

            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text.strip())
            return "\n".join(parts).strip()[:20000]
        except Exception:
            return raw.decode("utf-8", errors="ignore")[:20000]

    try:
        return raw.decode("utf-8", errors="ignore")[:20000]
    except Exception:
        return ""


_SKILL_PATTERNS = [
    r"\b(python|java|golang|go|rust|javascript|typescript|react|angular|vue)\b",
    r"\b(linux|rhel|ubuntu|centos|bash|shell)\b",
    r"\b(docker|kubernetes|k8s|helm|terraform|ansible|jenkins|ci/cd)\b",
    r"\b(aws|azure|gcp|cloud)\b",
    r"\b(mysql|postgres|mongodb|redis|kafka)\b",
    r"\b(nginx|apache|httpd|load balancer)\b",
    r"\b(prometheus|grafana|elk|splunk|datadog)\b",
    r"\b(networking|tcp/ip|dns|firewall|vpn)\b",
]


def parse_resume_text(text: str) -> dict:
    """Lightweight resume analysis without external LLM."""
    low = (text or "").lower()
    skills = []
    for pat in _SKILL_PATTERNS:
        skills.extend(re.findall(pat, low, flags=re.I))
    skills = sorted(set(s.lower() for s in skills))

    years = 0
    ym = re.search(r"(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?experience", low)
    if ym:
        years = int(ym.group(1))
    else:
        ym2 = re.search(r"experience[:\s]+(\d+)", low)
        if ym2:
            years = int(ym2.group(1))

    companies = re.findall(
        r"(?:at|@)\s+([A-Z][A-Za-z0-9&.\- ]{2,40})(?:\s|,|\||\n)",
        text or "",
    )[:5]

    roles = re.findall(
        r"(devops|sre|software engineer|system admin|linux admin|cloud engineer|platform engineer)",
        low,
    )

    return {
        "skills_detected": skills[:30],
        "years_experience_hint": years,
        "companies_mentioned": companies[:5],
        "roles_mentioned": list(dict.fromkeys(roles))[:5],
        "word_count": len((text or "").split()),
        "has_resume": bool(text and len(text.strip()) > 50),
    }


def build_profile_from_inputs(
    *,
    target_role: str = "",
    experience_level: str = "mid",
    years_experience: int = 0,
    current_company: str = "",
    secondary_technologies: list | None = None,
    primary_technology_name: str = "",
) -> dict:
    """Build resume_parsed hints from form selections when no resume file is uploaded."""
    skills = [s.lower() for s in (secondary_technologies or []) if s]
    if primary_technology_name:
        skills.insert(0, primary_technology_name.lower())
    roles = [target_role.lower()] if target_role else []
    if experience_level:
        roles.append(f"{experience_level} engineer")
    summary_parts = [p for p in [target_role, current_company, primary_technology_name] if p]
    synthetic_text = " ".join(summary_parts + skills)
    return {
        "skills_detected": sorted(set(skills))[:30],
        "years_experience_hint": years_experience or 0,
        "companies_mentioned": [current_company] if current_company else [],
        "roles_mentioned": roles[:5],
        "word_count": len(synthetic_text.split()),
        "has_resume": False,
        "source": "form_inputs",
    }


# ---------------------------------------------------------------------------
# Resume scoring (P2.5) — deterministic, 100% local, no external/LLM API.
# ---------------------------------------------------------------------------

# Per-technology keyword vocabulary used to score skills-match. Kept in sync
# with interview_ai._detect_topic so the resume score reflects the same
# vocabulary the live interviewer probes on. Keys are matched against the
# chosen technology name (and target role) by substring.
_TECH_KEYWORDS = {
    "kubernetes": ["kubernetes", "k8s", "kubectl", "pod", "deployment", "helm", "namespace", "ingress", "container"],
    "docker": ["docker", "container", "dockerfile", "image", "registry", "compose", "build"],
    "nginx": ["nginx", "reverse proxy", "upstream", "ssl", "load balanc", "tls", "proxy"],
    "linux": ["linux", "systemd", "kernel", "cgroup", "process", "bash", "shell", "rhel", "ubuntu"],
    "monitoring": ["prometheus", "grafana", "alertmanager", "metrics", "slo", "sli", "observability", "datadog"],
    "aws": ["aws", "ec2", "s3", "rds", "cloudwatch", "iam", "vpc", "lambda", "eks", "cloud"],
    "azure": ["azure", "aks", "blob", "cloud"],
    "gcp": ["gcp", "gke", "bigquery", "cloud"],
    "terraform": ["terraform", "tfstate", "provider", "module", "infrastructure as code", "iac"],
    "ansible": ["ansible", "playbook", "inventory", "role", "handler"],
    "ci_cd": ["ci/cd", "pipeline", "github actions", "jenkins", "gitlab", "argocd", "deploy"],
    "python": ["python", "django", "flask", "fastapi", "asyncio", "pip", "celery"],
    "security": ["security", "vulnerability", "cve", "secret", "credential", "rbac", "least privilege"],
    "database": ["database", "postgres", "mysql", "mongodb", "redis", "migration", "schema", "sql", "replication"],
    "networking": ["networking", "tcp/ip", "dns", "firewall", "vpn", "bgp", "routing"],
    "devops": ["ci/cd", "pipeline", "docker", "kubernetes", "terraform", "ansible", "automation", "infrastructure"],
    "sre": ["slo", "sli", "error budget", "incident", "on-call", "reliability", "mttr", "observability"],
}

# Action / impact verbs that signal an outcome-oriented, well-written resume.
_ACTION_VERBS = [
    "led", "built", "designed", "implemented", "migrated", "automated", "reduced",
    "improved", "scaled", "optimized", "deployed", "architected", "delivered",
    "launched", "managed", "owned", "resolved", "debugged", "increased", "decreased",
    "saved", "drove", "shipped", "mentored", "established", "introduced",
]

# Section headings a strong, well-structured resume usually contains.
_STRUCTURE_SECTIONS = [
    "experience", "skills", "education", "project", "summary", "certification",
    "achievement", "work history", "professional",
]

# Rough years-of-experience expectation per chosen level, for the experience subscore.
_LEVEL_YEARS = {"junior": 1, "mid": 4, "senior": 8, "lead": 12}

_QUANT_RE = re.compile(
    r"(\d+\s*%|\$\s*\d|\d+\s*(?:x|k|m|gb|tb|ms|rps|qps|req|users?|hours?|hrs?|days?|weeks?|months?|servers?|nodes?|pods?|instances?))",
    re.I,
)


def _tech_vocabulary(target_technology: str, target_role: str) -> tuple[str, list[str]]:
    """Resolve the chosen technology/role to a keyword vocabulary.

    Returns (matched_label, keywords). Falls back to a broad devops/SRE
    vocabulary when nothing matches so the score is never undefined.
    """
    hay = f"{target_technology or ''} {target_role or ''}".lower()
    matched: list[str] = []
    label = ""
    for tech, kws in _TECH_KEYWORDS.items():
        token = tech.replace("_", " ")
        if token in hay or any(k in hay for k in (tech, token)):
            matched.extend(kws)
            label = label or tech
    if not matched:
        # Generic infra/devops baseline keeps scoring meaningful for any role.
        matched = sorted({k for kws in (_TECH_KEYWORDS["devops"], _TECH_KEYWORDS["linux"]) for k in kws})
        label = "general"
    return label, sorted(set(matched))


def score_resume(
    parsed: dict | None,
    *,
    resume_text: str = "",
    target_technology: str = "",
    target_role: str = "",
    experience_level: str = "mid",
    years_experience: int = 0,
) -> dict:
    """Score a resume 0–100 against the chosen technology / level / role.

    Deterministic and fully local (no LLM / external API). Combines four
    weighted subscores:

    - ``skills_match``  (35%): coverage of the chosen technology's keyword set.
    - ``experience``    (25%): claimed years vs the level's expectation.
    - ``clarity``       (20%): length, structure (sections), action verbs.
    - ``keywords``      (20%): quantified impact + breadth of detected skills.

    Returns ``overall_score``, the four ``subscores``, the ``vocabulary`` used,
    and concrete ``tips`` to improve weak areas.
    """
    parsed = parsed or {}
    text = (resume_text or "").strip()
    low = text.lower()
    has_resume = bool(parsed.get("has_resume")) or len(low) > 50

    label, vocab = _tech_vocabulary(target_technology, target_role)

    if not has_resume:
        return {
            "overall_score": None,
            "subscores": {},
            "matched_keywords": [],
            "missing_keywords": [],
            "vocabulary": label,
            "tips": [
                "Upload a resume (PDF or DOCX) to see your resume score and get tailored improvement tips.",
            ],
            "has_resume": False,
            "message": "No resume uploaded",
        }

    detected = [s.lower() for s in (parsed.get("skills_detected") or [])]

    # --- 1. Skills match: how many of the chosen tech's keywords appear. ---
    pool = (low + " " + " ".join(detected)).strip()
    matched_kws = sorted({k for k in vocab if k in pool})
    missing_kws = [k for k in vocab if k not in pool]
    skills_match = round(100 * len(matched_kws) / len(vocab)) if vocab else 0
    # Soft floor so a near-empty resume still gets a small, honest signal.
    skills_match = max(0, min(100, skills_match))

    # --- 2. Experience: claimed years vs the level expectation. ---
    expected_years = _LEVEL_YEARS.get((experience_level or "mid").lower(), 4)
    claimed_years = int(parsed.get("years_experience_hint") or years_experience or 0)
    if expected_years <= 0:
        experience = 70
    else:
        ratio = claimed_years / expected_years
        experience = round(100 * min(1.0, ratio)) if ratio >= 0.5 else round(60 * ratio)
    experience = max(0, min(100, experience))

    # --- 3. Clarity / structure: length, sections, action verbs. ---
    word_count = int(parsed.get("word_count") or len(low.split()))
    if word_count <= 0:
        length_score = 0
    elif word_count < 150:
        length_score = round(60 * word_count / 150)
    elif word_count <= 900:
        length_score = 100
    elif word_count <= 1400:
        length_score = 80
    else:
        length_score = 60
    sections_found = sum(1 for s in _STRUCTURE_SECTIONS if s in low)
    structure_score = min(100, sections_found * 25)
    verbs_found = sorted({v for v in _ACTION_VERBS if re.search(rf"\b{v}\b", low)})
    verb_score = min(100, len(verbs_found) * 12)
    clarity = round(0.4 * length_score + 0.3 * structure_score + 0.3 * verb_score)
    clarity = max(0, min(100, clarity))

    # --- 4. Keywords: quantified impact + breadth of detected skills. ---
    quant_hits = len(_QUANT_RE.findall(text))
    quant_score = min(100, quant_hits * 20)
    breadth_score = min(100, len(detected) * 12)
    keywords = round(0.55 * quant_score + 0.45 * breadth_score)
    keywords = max(0, min(100, keywords))

    subscores = {
        "skills_match": skills_match,
        "experience": experience,
        "clarity": clarity,
        "keywords": keywords,
    }
    overall = round(
        0.35 * skills_match + 0.25 * experience + 0.20 * clarity + 0.20 * keywords
    )
    overall = max(0, min(100, overall))

    # --- Concrete, prioritized improvement tips (weakest areas first). ---
    tips: list[str] = []
    if not has_resume:
        if len(low) > 0:
            tips.append(
                "This resume is too short to analyze — add real experience, skills, "
                "and project detail so we can score it properly."
            )
        else:
            tips.append(
                "Upload an actual resume (PDF or DOCX) — we're scoring from your form "
                "inputs only, so the score is a rough estimate."
            )
    if skills_match < 70 and missing_kws:
        show = ", ".join(missing_kws[:6])
        tips.append(
            f"Add concrete {label if label != 'general' else 'role-relevant'} "
            f"keywords the role expects: {show}."
        )
    if experience < 60:
        tips.append(
            f"State total years of experience clearly (e.g. \"{expected_years}+ years\") — "
            f"a {experience_level}-level role expects around {expected_years} years."
        )
    if quant_hits < 3:
        tips.append(
            "Quantify your impact — add numbers like \"cut MTTR 40%\", \"saved $20k/mo\", "
            "or \"scaled to 2M req/day\". Resumes with metrics score higher."
        )
    if len(verbs_found) < 5:
        tips.append(
            "Start bullet points with strong action verbs (Led, Built, Automated, "
            "Reduced) instead of \"Responsible for\"."
        )
    if sections_found < 3:
        tips.append(
            "Use clear section headings (Summary, Experience, Skills, Projects, "
            "Education) so the structure is easy to scan."
        )
    if word_count and word_count < 150:
        tips.append("Your resume looks short — expand on responsibilities, scope, and outcomes.")
    elif word_count > 1400:
        tips.append("Your resume is long — tighten it to the most relevant, recent, high-impact work.")
    if not tips:
        tips.append("Strong resume for this role — keep tailoring keywords to each specific job description.")

    return {
        "overall_score": overall,
        "subscores": subscores,
        "matched_keywords": matched_kws[:20],
        "missing_keywords": missing_kws[:12],
        "vocabulary": label,
        "tips": tips[:6],
        "has_resume": has_resume,
    }
