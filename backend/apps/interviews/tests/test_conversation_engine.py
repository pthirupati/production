"""Tests for the free offline conversational interview engine (Phase 0/1)."""

from __future__ import annotations

import io
import random
import re
from pathlib import Path

from django.test import TestCase

from apps.interviews.services.conversation import (
    analyze_answer,
    decide_next_move,
    generate_follow_up_question,
    normalize_transcript,
)
from apps.interviews.services.conversation.memory import CampaignMemory
from apps.interviews.services.conversation.policy import NextMove
from apps.interviews.services.conversation.scorer import compute_semantic_scores
from apps.interviews.services.conversation_intelligence import detect_contradiction
from apps.interviews.services.resume_parser import extract_text_from_upload
from apps.interviews.services.scoring import aggregate_round_scores


class ContradictionDetectionTest(TestCase):
    def test_opposite_pair_triggers_consistency_probe(self):
        memory = {"claims": ["We must use rolling deploys for production releases"]}
        new_answer = "Actually we avoid rolling deploys for production releases — blue-green only."
        prior = detect_contradiction(memory, new_answer)
        self.assertIsNotNone(prior, "planted contradiction should be detected")
        self.assertIn("rolling", prior.lower())


class SttRepairTest(TestCase):
    def test_kube_cuttle_becomes_kubectl(self):
        self.assertIn("kubectl", normalize_transcript("I ran kube cuttle get pods"))

    def test_cooper_neties_becomes_kubernetes(self):
        self.assertIn("kubernetes", normalize_transcript("we use cooper neties in prod"))

    def test_no_js_becomes_nodejs(self):
        self.assertIn("Node.js", normalize_transcript("built with no js backend"))


class NextMovePolicyTest(TestCase):
    def _decide(self, answer: str, *, memory: CampaignMemory | None = None, **kwargs):
        analysis = analyze_answer(answer_text=answer, question_text="How do you debug pods?")
        return decide_next_move(
            analysis=analysis,
            memory=memory or CampaignMemory(),
            **kwargs,
        )

    def test_vague_answer_triggers_clarify(self):
        d = self._decide(
            "I think we probably checked logs and metrics but the details are fuzzy on the exact steps.",
        )
        self.assertEqual(d.move, NextMove.CLARIFY)

    def test_idk_triggers_hint_then_move(self):
        d = self._decide("I don't know.")
        self.assertEqual(d.move, NextMove.HINT_THEN_MOVE)

    def test_candidate_question_triggers_answer_candidate(self):
        d = self._decide("What kind of scenarios will we cover today?")
        self.assertEqual(d.move, NextMove.ANSWER_CANDIDATE)

    def test_contradiction_triggers_challenge(self):
        mem = CampaignMemory(claims=["We must use rolling deploys for production"])
        d = self._decide(
            "We avoid rolling deploys in production — only blue-green.",
            memory=mem,
        )
        self.assertEqual(d.move, NextMove.CHALLENGE)

    def test_strong_depth_triggers_drill_down(self):
        answer = (
            "I ran kubectl describe pod api-7f2 and saw OOMKilled at 512Mi. "
            "Then kubectl top pod showed memory at 98% for 20 minutes before restart."
        )
        d = self._decide(answer)
        self.assertIn(d.move, (NextMove.DRILL_DOWN, NextMove.SCENARIO_ESCALATE))

    def test_nervous_brief_streak_triggers_ease_redirect(self):
        d = self._decide("Um, maybe logs?", brief_streak=2)
        self.assertEqual(d.move, NextMove.EASE_REDIRECT)

    def test_strong_streak_triggers_scenario_escalate(self):
        answer = (
            "I used kubectl rollout status, verified readiness probes returned 200, "
            "and confirmed error rate dropped from 12% to 0.1% in 3 minutes."
        )
        mem = CampaignMemory(competence_estimate=0.7)
        d = self._decide(answer, memory=mem, strong_streak=3)
        self.assertEqual(d.move, NextMove.SCENARIO_ESCALATE)


class GroundedFollowUpTest(TestCase):
    def test_follow_up_contains_candidate_token(self):
        answer = (
            "I ran kubectl rollout status deployment/api, then kubectl top pod "
            "and saw memory at 98% for 20 minutes before the restart."
        )
        analysis = analyze_answer(answer_text=answer, question_text="How do you deploy?")
        decision = decide_next_move(analysis=analysis, memory=CampaignMemory())
        self.assertNotEqual(decision.move, NextMove.NEW_TOPIC)
        rng = random.Random(42)
        _, question = generate_follow_up_question(
            analysis=analysis,
            decision=decision,
            used_texts=set(),
            rng=rng,
        )
        self.assertTrue(question)
        low = question.lower()
        self.assertTrue(
            "kubectl" in low or "deployment" in low or "rollout" in low,
            f"follow-up should reference candidate content, got: {question!r}",
        )


class SemanticScorerTest(TestCase):
    def test_concise_correct_beats_long_irrelevant(self):
        concise = compute_semantic_scores(
            candidate_answer="Use kubectl describe pod and check OOMKilled events; then raise limits or fix the leak.",
            question_text="How do you debug a CrashLoopBackOff pod?",
            round_type="technical",
            expected_keywords=["kubectl", "describe", "logs"],
        )
        ramble = compute_semantic_scores(
            candidate_answer=(
                "Synergy stakeholder paradigm leverage blockchain agile scrum waterfall "
                "in my opinion basically literally you know um uh synergy synergy synergy "
                "stakeholder pivot moonshot disrupt unicorn bandwidth paradigm shift "
            ) * 4,
            question_text="How do you debug a CrashLoopBackOff pod?",
            round_type="technical",
            expected_keywords=["kubectl", "describe", "logs"],
        )
        self.assertGreater(
            concise["composite_score"],
            ramble["composite_score"],
            "concise relevant answer must outscore jargon wall",
        )

    def test_disfluent_correct_not_heavily_penalized(self):
        clean = compute_semantic_scores(
            candidate_answer="I run kubectl describe pod and check events for OOMKilled.",
            question_text="Debug CrashLoopBackOff?",
            round_type="technical",
            expected_keywords=["kubectl"],
        )
        disfluent = compute_semantic_scores(
            candidate_answer="Um, uh, I like, run kubectl describe pod and, you know, check events for OOMKilled.",
            question_text="Debug CrashLoopBackOff?",
            round_type="technical",
            expected_keywords=["kubectl"],
        )
        self.assertGreaterEqual(disfluent["composite_score"], clean["composite_score"] * 0.75)


class IndependentDimensionScoringTest(TestCase):
    def test_presence_and_resume_not_hardcoded(self):
        rows = [
            {
                "content": "I used kubectl and prometheus to cut MTTR from 45m to 8m.",
                "score": 82,
                "metadata": {
                    "quality": "strong",
                    "depth_score": 70,
                    "concrete_score": 75,
                    "star_score": 80,
                    "relevance_score": 78,
                    "keyword_hit_rate": 0.6,
                },
            },
            {
                "content": "um uh like you know basically",
                "score": 30,
                "metadata": {"quality": "weak", "user_skip": False},
            },
        ]
        agg = aggregate_round_scores([82, 30], round_type="technical", answer_rows=rows)
        self.assertNotEqual(agg["presence_score"], 72.0)
        self.assertNotEqual(agg["resume_alignment_score"], 68.0)
        self.assertNotEqual(agg["technical_score"], agg["communication_score"])


class DocxParsingTest(TestCase):
    def test_docx_extracts_paragraph_text(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Senior DevOps Engineer with Kubernetes experience.")
        doc.add_paragraph("Skills: Docker, Terraform, Prometheus.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        class FakeUpload:
            name = "resume.docx"

            def read(self):
                return buf.getvalue()

            def seek(self, pos):
                buf.seek(pos)

        text = extract_text_from_upload(FakeUpload())
        self.assertIn("Kubernetes", text)
        self.assertIn("Terraform", text)


class NoPaidSdkImportsTest(TestCase):
    """CI assert: nothing under apps/interviews imports paid SDKs."""

    FORBIDDEN = (
        "import anthropic",
        "from anthropic",
        "import openai",
        "from openai",
        "import elevenlabs",
        "from elevenlabs",
    )

    def test_no_paid_sdk_imports_under_interviews_app(self):
        root = Path(__file__).resolve().parents[1]
        offenders: list[str] = []
        for path in root.rglob("*.py"):
            if path.name.startswith("test_"):
                continue
            src = path.read_text(encoding="utf-8", errors="ignore")
            code = re.sub(r'""".*?"""', "", src, flags=re.DOTALL)
            code = re.sub(r"'''.*?'''", "", code, flags=re.DOTALL)
            code = "\n".join(line.split("#", 1)[0] for line in code.splitlines()).lower()
            for forbidden in self.FORBIDDEN:
                if forbidden in code:
                    offenders.append(f"{path.relative_to(root)}: {forbidden}")
        self.assertEqual(offenders, [], f"paid SDK imports found: {offenders}")
